"""Correction Document Renderer.

Produces student-facing correction documents in DOCX and HTML formats.

Each MC/MA item with per-choice rationales is rendered as a 4-column table:

    Letter | Choice Text | ✓ / ✗ | Rationale

Formatting spec source: dev/QuizForge_PerChoiceRationale_Spec.md, Sections 2–3.

Usage::

    from engine.rendering.correction_doc.renderer import CorrectionDocRenderer
    from engine.spec_engine.packager import package_quiz

    packaged = package_quiz(quiz_payload)
    renderer = CorrectionDocRenderer()

    html_str  = renderer.render_html(packaged)
    docx_bytes = renderer.render_docx(packaged)
"""

from __future__ import annotations

import io
from typing import Dict, List, Optional, Tuple

from ...spec_engine.models import ChoiceRationale, PackagedQuiz, RationalesEntry

# ---------------------------------------------------------------------------
# Colour constants — Section 3 of the spec
# ---------------------------------------------------------------------------

HEADER_BG_HEX = "1F3864"        # dark blue — table header row background
HEADER_FG_HEX = "FFFFFF"        # white    — table header row text
CORRECT_ROW_BG_HEX = "E6F3E6"   # light green — correct-answer row background
CORRECT_MARK_HEX = "006600"     # dark green — ✓ symbol colour & correct text
INCORRECT_ROW_BG_HEX = "F3E6E6" # light red — incorrect-answer row background
INCORRECT_MARK_HEX = "CC0000"   # dark red — ✗ symbol colour & incorrect text
QUESTION_HDR_HEX = "1F3864"     # dark blue — question stem paragraph

# Fractional column widths (as proportions of usable page width).
# Letter : Choice : Mark : Rationale
_COL_RATIOS = (0.07, 0.38, 0.07, 0.48)

# Usable page width in inches (Letter paper minus 0.75" margins each side)
_PAGE_WIDTH_INCHES = 7.0

# Approximate word-count guard for "no rationale available" fallback text
_FALLBACK_RATIONALE = "(no rationale provided)"


# ---------------------------------------------------------------------------
# Internal helpers — shared between both renderers
# ---------------------------------------------------------------------------

def _build_item_index(items: List[Dict]) -> Dict[str, Dict]:
    """Return a dict keyed by item id for fast lookup."""
    return {item["id"]: item for item in items if isinstance(item.get("id"), str)}


def _choice_letter(choice: Dict, index: int) -> str:
    """Return the display letter for a choice (its id field, or A/B/C/D by position)."""
    cid = choice.get("id")
    if isinstance(cid, str) and cid:
        return cid.upper()
    return chr(65 + index)


def _build_rationale_index(rationales: List[RationalesEntry]) -> Dict[str, RationalesEntry]:
    return {r.item_id: r for r in rationales}


def _strip_html(text: str) -> str:
    """Strip HTML tags and unescape entities, collapsing runs of whitespace."""
    import html
    import re
    # Replace block-level tags with a space so words from adjacent paragraphs don't run together
    text = re.sub(r"</(p|div|br|li|h[1-6])[^>]*>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>" , " ", text, flags=re.IGNORECASE)
    # Strip all remaining tags
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    # Collapse multiple spaces/newlines
    return re.sub(r"\s+", " ", text).strip()


def _items_to_render(
    items: List[Dict],
    rationale_index: Dict[str, RationalesEntry],
) -> List[Tuple[int, Dict, RationalesEntry]]:
    """Return (display_number, item, rationale_entry) for every renderable item.

    Only MC and MA items that have a matching per-choice rationale entry are included.
    Items without a valid rationale entry are silently omitted.
    Returns items in their original order; question numbering skips stimuli.
    """
    result = []
    q_number = 1
    for item in items:
        qtype = item.get("type", "")
        if qtype in ("STIMULUS", "STIMULUS_END"):
            continue
        if qtype in ("MC", "MA"):
            item_id = item.get("id")
            if item_id and item_id in rationale_index:
                result.append((q_number, item, rationale_index[item_id]))
        q_number += 1
    return result


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

def _shade_cell(cell, fill_hex: str) -> None:
    """Apply background fill colour to a python-docx table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_col_widths(table, col_widths_inches: Tuple[float, ...]) -> None:
    """Set absolute column widths on a python-docx table."""
    from docx.shared import Inches

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_inches):
                cell.width = Inches(col_widths_inches[i])


def _render_docx_item(
    doc,
    q_number: int,
    item: Dict,
    entry: RationalesEntry,
) -> None:
    """Add one question block (header + table) to a python-docx Document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    # --- Question header ---
    hdr = doc.add_paragraph()
    run = hdr.add_run(f"Q{q_number}: {_strip_html(item.get('prompt', ''))}")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(QUESTION_HDR_HEX)
    run.font.size = Pt(11)

    # --- Column widths ---
    col_widths = tuple(r * _PAGE_WIDTH_INCHES for r in _COL_RATIOS)

    # Per-choice table
    if entry.choices:
        choices_by_id = {c.id: c for c in entry.choices}
        item_choices = item.get("choices", [])

        table = doc.add_table(rows=1 + len(item_choices), cols=4)
        table.style = "Table Grid"

        # Header row
        hdr_row = table.rows[0]
        for cell, label in zip(hdr_row.cells, ["", "Choice", "✓ / ✗", "Rationale"]):
            _shade_cell(cell, HEADER_BG_HEX)
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(HEADER_FG_HEX)
            run.font.size = Pt(10)

        # Data rows
        for row_idx, (choice_dict, data_row) in enumerate(
            zip(item_choices, table.rows[1:])
        ):
            letter = _choice_letter(choice_dict, row_idx)
            is_correct = bool(choice_dict.get("correct"))
            choice_text = choice_dict.get("text", "")
            cr: Optional[ChoiceRationale] = choices_by_id.get(letter)
            rationale_text = cr.rationale if cr else _FALLBACK_RATIONALE

            row_bg = CORRECT_ROW_BG_HEX if is_correct else INCORRECT_ROW_BG_HEX
            text_color = CORRECT_MARK_HEX if is_correct else INCORRECT_MARK_HEX

            # Apply row shading and populate cells
            cells = data_row.cells

            # Col 0 — letter
            _shade_cell(cells[0], row_bg)
            p0 = cells[0].paragraphs[0]
            r0 = p0.add_run(letter)
            r0.bold = True
            r0.font.color.rgb = RGBColor.from_string(text_color)
            r0.font.size = Pt(10)

            # Col 1 — choice text
            _shade_cell(cells[1], row_bg)
            p1 = cells[1].paragraphs[0]
            r1 = p1.add_run(choice_text)
            r1.bold = is_correct
            r1.font.color.rgb = RGBColor.from_string(text_color)
            r1.font.size = Pt(10)

            # Col 2 — ✓ / ✗ (centered)
            _shade_cell(cells[2], row_bg)
            p2 = cells[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mark = "✓" if is_correct else "✗"
            mark_color = CORRECT_MARK_HEX if is_correct else INCORRECT_MARK_HEX
            r2 = p2.add_run(mark)
            r2.bold = True
            r2.font.color.rgb = RGBColor.from_string(mark_color)
            r2.font.size = Pt(10)

            # Col 3 — rationale
            _shade_cell(cells[3], row_bg)
            p3 = cells[3].paragraphs[0]
            r3 = p3.add_run(rationale_text)
            r3.italic = not is_correct
            r3.font.color.rgb = RGBColor.from_string(text_color)
            r3.font.size = Pt(10)

        _set_col_widths(table, col_widths)

    # Spacer after each item
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# HTML renderer
# ---------------------------------------------------------------------------

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} — Correction Document</title>
<style>
  body {{
    font-family: Calibri, 'Segoe UI', Arial, sans-serif;
    font-size: 11pt;
    color: #111;
    margin: 1in 0.75in;
    line-height: 1.4;
  }}
  h1.doc-title {{
    color: #{header_bg};
    font-size: 15pt;
    margin-bottom: 0.25em;
  }}
  p.doc-subtitle {{
    font-size: 10pt;
    color: #555;
    margin-top: 0;
    margin-bottom: 1.5em;
  }}
  .question-block {{
    margin-bottom: 1.5em;
  }}
  .question-header {{
    font-weight: bold;
    color: #{q_hdr};
    font-size: 11pt;
    margin-bottom: 0.4em;
  }}
  table.rationale-table {{
    border-collapse: collapse;
    width: 100%;
    font-size: 10pt;
  }}
  table.rationale-table th,
  table.rationale-table td {{
    border: 1px solid #999;
    padding: 5px 7px;
    vertical-align: top;
  }}
  table.rationale-table thead tr {{
    background-color: #{header_bg};
    color: #{header_fg};
    font-weight: bold;
  }}
  table.rationale-table thead th {{
    border-color: #{header_bg};
  }}
  .col-letter   {{ width: 4%;  text-align: center; font-weight: bold; }}
  .col-choice   {{ width: 36%; }}
  .col-mark     {{ width: 5%;  text-align: center; font-weight: bold; }}
  .col-rationale{{ width: 55%; }}
  .correct-row  {{ background-color: #{correct_bg}; }}
  .incorrect-row {{ background-color: #{incorrect_bg}; }}
  .correct-mark {{ color: #{correct_mark}; }}
  .incorrect-mark {{ color: #{incorrect_mark}; }}
  .correct-choice {{ font-weight: bold; color: #{correct_mark}; }}
  .incorrect-choice {{ color: #{incorrect_mark}; }}
  .correct-text {{ color: #{correct_mark}; }}
  .incorrect-text {{ color: #{incorrect_mark}; font-style: italic; }}
</style>
</head>
<body>
<h1 class="doc-title">{title}</h1>
<p class="doc-subtitle">Correction Document — review each answer choice and its explanation.</p>
{body}
</body>
</html>
"""


def _render_html_item(q_number: int, item: Dict, entry: RationalesEntry) -> str:
    """Return the HTML for a single question block."""
    prompt = item.get("prompt", "")
    header = f'<div class="question-header">Q{q_number}: {_escape_html(prompt)}</div>'

    table = ""
    if entry.choices:
        choices_by_id = {c.id: c for c in entry.choices}
        item_choices = item.get("choices", [])

        rows_html = []
        for row_idx, choice_dict in enumerate(item_choices):
            letter = _choice_letter(choice_dict, row_idx)
            is_correct = bool(choice_dict.get("correct"))
            choice_text = choice_dict.get("text", "")
            cr: Optional[ChoiceRationale] = choices_by_id.get(letter)
            rationale_text = cr.rationale if cr else _FALLBACK_RATIONALE

            row_class = ' class="correct-row"' if is_correct else ' class="incorrect-row"'
            mark = "✓" if is_correct else "✗"
            mark_class = "correct-mark" if is_correct else "incorrect-mark"
            choice_class = "correct-choice" if is_correct else "incorrect-choice"
            rationale_class = "correct-text" if is_correct else "incorrect-text"

            rows_html.append(
                f'<tr{row_class}>'
                f'<td class="col-letter">{_escape_html(letter)}</td>'
                f'<td class="col-choice {choice_class}">{_escape_html(choice_text)}</td>'
                f'<td class="col-mark"><span class="{mark_class}">{mark}</span></td>'
                f'<td class="col-rationale {rationale_class}">{_escape_html(rationale_text)}</td>'
                f"</tr>"
            )

        table = (
            '<table class="rationale-table">'
            "<thead><tr>"
            '<th class="col-letter"></th>'
            '<th class="col-choice">Choice</th>'
            '<th class="col-mark">✓ / ✗</th>'
            '<th class="col-rationale">Rationale</th>'
            "</tr></thead>"
            "<tbody>"
            + "\n".join(rows_html)
            + "</tbody></table>"
        )

    return f'<div class="question-block">{header}{table}</div>'


def _escape_html(text: str) -> str:
    """Minimal HTML escaping for text content."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ---------------------------------------------------------------------------
# Public renderer class
# ---------------------------------------------------------------------------

class CorrectionDocRenderer:
    """Render per-choice rationale correction documents.

    Accepts a :class:`~engine.spec_engine.models.PackagedQuiz` and produces:
    - A DOCX file (bytes) formatted for printing
    - A self-contained HTML string for Canvas or browser viewing

    Only MC/MA items with a per-choice rationale entry (``choices`` array) are
    rendered. Items with no rationale entry are silently omitted.
    """

    def render_docx(self, quiz: PackagedQuiz) -> bytes:
        """Return DOCX bytes for the correction document.

        Requires ``python-docx`` (already in requirements.txt).

        Args:
            quiz: A packaged quiz produced by :func:`engine.spec_engine.packager.package_quiz`.

        Returns:
            Raw bytes of the generated ``.docx`` file.
        """
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Document title
        title_text = quiz.title or "Correction Document"
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{title_text} — Correction Document")
        title_run.bold = True
        title_run.font.size = Pt(15)
        from docx.shared import RGBColor
        title_run.font.color.rgb = RGBColor.from_string(HEADER_BG_HEX)

        subtitle_para = doc.add_paragraph(
            "Review each answer choice and its explanation."
        )
        subtitle_para.runs[0].font.size = Pt(10)

        rationale_index = _build_rationale_index(quiz.rationales)
        renderable = _items_to_render(quiz.items, rationale_index)

        if not renderable:
            doc.add_paragraph(
                "No per-choice rationales found. Check that the quiz was generated "
                "with the per-choice rationale format."
            )
        else:
            for q_number, item, entry in renderable:
                _render_docx_item(doc, q_number, item, entry)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_html(self, quiz: PackagedQuiz) -> str:
        """Return a self-contained HTML string for the correction document.

        The returned string is a complete ``<!DOCTYPE html>`` document with
        inline CSS — suitable for saving as ``.html`` or embedding in Canvas.

        Args:
            quiz: A packaged quiz produced by :func:`engine.spec_engine.packager.package_quiz`.

        Returns:
            UTF-8 string containing the full HTML document.
        """
        rationale_index = _build_rationale_index(quiz.rationales)
        renderable = _items_to_render(quiz.items, rationale_index)

        if renderable:
            body = "\n".join(
                _render_html_item(q_number, item, entry)
                for q_number, item, entry in renderable
            )
        else:
            body = (
                "<p><em>No per-choice rationales found. Check that the quiz was "
                "generated with the per-choice rationale format.</em></p>"
            )

        title = _escape_html(quiz.title or "Correction Document")
        return _HTML_TEMPLATE.format(
            title=title,
            body=body,
            header_bg=HEADER_BG_HEX,
            header_fg=HEADER_FG_HEX,
            q_hdr=QUESTION_HDR_HEX,
            correct_bg=CORRECT_ROW_BG_HEX,
            incorrect_bg=INCORRECT_ROW_BG_HEX,
            correct_mark=CORRECT_MARK_HEX,
            incorrect_mark=INCORRECT_MARK_HEX,
        )
