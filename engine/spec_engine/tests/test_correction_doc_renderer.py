"""Tests for the CorrectionDocRenderer (HTML and DOCX output).

These tests operate on the spec_engine's PackagedQuiz / RationalesEntry model.
The DOCX tests verify that valid bytes are produced and that python-docx can
re-parse them; they do not attempt to inspect raw XML unless specifically needed.
The HTML tests verify the semantic content and key formatting markers.
"""

from __future__ import annotations

import json
import re

import pytest

from ..parser import parse_news_json
from ..packager import package_quiz
from ...rendering.correction_doc.renderer import (
    CorrectionDocRenderer,
    CORRECT_MARK_HEX,
    CORRECT_ROW_BG_HEX,
    HEADER_BG_HEX,
    INCORRECT_MARK_HEX,
)


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------

def _wrap(payload: dict) -> str:
    return f"<QUIZFORGE_JSON>{json.dumps(payload)}</QUIZFORGE_JSON>"


def _packaged_quiz(raw_payload: dict):
    return package_quiz(parse_news_json(_wrap(raw_payload)))


def _per_choice_payload():
    return {
        "version": "3.0-json",
        "title": "Flowchart Quiz",
        "items": [
            {
                "id": "q1",
                "type": "MC",
                "prompt": "What does an arrow labeled 'Yes' represent in a flowchart?",
                "choices": [
                    {"id": "A", "text": "The path taken when the decision is Yes", "correct": True},
                    {"id": "B", "text": "The value stored in a variable", "correct": False},
                    {"id": "C", "text": "The number of loop iterations", "correct": False},
                    {"id": "D", "text": "The data type of the result", "correct": False},
                ],
            },
            {
                "id": "q2",
                "type": "MC",
                "prompt": "Which shape represents a decision in a flowchart?",
                "choices": [
                    {"id": "A", "text": "Rectangle", "correct": False},
                    {"id": "B", "text": "Diamond", "correct": True},
                    {"id": "C", "text": "Oval", "correct": False},
                ],
            },
        ],
        "rationales": [
            {
                "item_id": "q1",
                "choices": [
                    {"id": "A", "correct": True,  "rationale": "Arrows show program flow; 'Yes' marks the path when the decision is true."},
                    {"id": "B", "correct": False, "rationale": "Variables store data values; they are unrelated to arrows."},
                    {"id": "C", "correct": False, "rationale": "Loop counts are not shown by individual arrows."},
                    {"id": "D", "correct": False, "rationale": "Data types describe variable kinds, not arrow labels."},
                ],
            },
            {
                "item_id": "q2",
                "choices": [
                    {"id": "A", "correct": False, "rationale": "Rectangles represent process steps, not decisions."},
                    {"id": "B", "correct": True,  "rationale": "Diamonds represent decision points with Yes/No branches."},
                    {"id": "C", "correct": False, "rationale": "Ovals mark the start and end of a flowchart."},
                ],
            },
        ],
    }


def _non_mc_payload():
    """Payload where the only rationale is for a TF item (should not render a table)."""
    return {
        "version": "3.0-json",
        "title": "TF Quiz",
        "items": [
            {
                "id": "q1",
                "type": "TF",
                "prompt": "The Earth is round.",
                "answer": True,
            }
        ],
        "rationales": [
            {"item_id": "q1", "rationale": "The Earth is an oblate spheroid."}
        ],
    }


def _mixed_types_payload():
    """Quiz with a STIMULUS, TF, and MC — only the MC should be rendered."""
    return {
        "version": "3.0-json",
        "title": "Mixed Types",
        "items": [
            {"id": "s1", "type": "STIMULUS", "prompt": "Read the following."},
            {"id": "q1", "type": "TF", "prompt": "True or false?", "answer": True},
            {
                "id": "q2",
                "type": "MC",
                "prompt": "MC question here",
                "choices": [
                    {"id": "A", "text": "Choice A", "correct": True},
                    {"id": "B", "text": "Choice B", "correct": False},
                ],
            },
            {"id": "se1", "type": "STIMULUS_END", "prompt": ""},
        ],
        "rationales": [
            {"item_id": "q1", "rationale": "It is true."},
            {
                "item_id": "q2",
                "choices": [
                    {"id": "A", "correct": True,  "rationale": "A is correct."},
                    {"id": "B", "correct": False, "rationale": "B describes something else."},
                ],
            },
        ],
    }


# ---------------------------------------------------------------------------
# HTML output tests
# ---------------------------------------------------------------------------

class TestHtmlOutput:

    def test_html_is_valid_document(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert html.startswith("<!DOCTYPE html>")
        assert "<html" in html
        assert "</html>" in html

    def test_html_contains_title(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "Flowchart Quiz" in html

    def test_html_contains_question_stems(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "labeled &#x27;Yes&#x27;" in html or "labeled 'Yes'" in html or "labeled" in html
        assert "Which shape represents a decision" in html

    def test_html_contains_choice_text(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "The path taken when the decision is Yes" in html
        assert "The value stored in a variable" in html
        assert "Diamond" in html

    def test_html_contains_rationale_text(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "Arrows show program flow" in html
        assert "Variables store data values" in html
        assert "Diamonds represent decision points" in html

    def test_html_correct_answer_check_mark(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "✓" in html

    def test_html_incorrect_answer_cross_mark(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "✗" in html

    def test_html_correct_row_background_color(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        # The CSS class or the hex colour for the correct row must be present
        assert CORRECT_ROW_BG_HEX in html or "correct-row" in html

    def test_html_correct_mark_color_in_css(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert CORRECT_MARK_HEX in html

    def test_html_incorrect_mark_color_in_css(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert INCORRECT_MARK_HEX in html

    def test_html_header_color_in_css(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert HEADER_BG_HEX in html

    def test_html_incorrect_rationale_italic_class(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "incorrect-rationale" in html

    def test_html_correct_choice_bold_class(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "correct-choice" in html

    def test_html_question_numbering(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "Q1:" in html
        assert "Q2:" in html

    def test_html_two_questions_two_tables(self):
        quiz = _packaged_quiz(_per_choice_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert html.count('<table class="rationale-table">') == 2

    def test_html_non_mc_items_skipped(self):
        """TF items should not produce correction tables even with rationales."""
        quiz = _packaged_quiz(_non_mc_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        # No MC items → should show the "no rationales found" message
        assert "No per-choice rationales found" in html or "rationale-table" not in html

    def test_html_only_mc_rendered_in_mixed_quiz(self):
        """STIMULUS and TF items must not appear as table blocks."""
        quiz = _packaged_quiz(_mixed_types_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        assert "MC question here" in html
        assert '<table class="rationale-table">' in html
        # The TF prompt should not generate a table — only the MC item should
        assert html.count('<table class="rationale-table">') == 1

    def test_html_question_number_respects_non_mc_items(self):
        """The TF item counts toward the question number; MC after it is Q2."""
        quiz = _packaged_quiz(_mixed_types_payload())
        html = CorrectionDocRenderer().render_html(quiz)
        # q2 is the 2nd scored item (after TF q1), so it should be labelled Q2
        assert "Q2:" in html

    def test_html_no_rationale_empty_message(self):
        payload = {
            "version": "3.0-json",
            "title": "Empty",
            "items": [
                {"id": "q1", "type": "MC", "prompt": "Q?",
                 "choices": [{"id": "A", "text": "A", "correct": True}]}
            ],
            "rationales": [],
        }
        quiz = _packaged_quiz(payload)
        html = CorrectionDocRenderer().render_html(quiz)
        assert "No per-choice rationales found" in html


# ---------------------------------------------------------------------------
# DOCX output tests
# ---------------------------------------------------------------------------

class TestDocxOutput:

    def test_docx_returns_bytes(self):
        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_docx_is_valid_zip_magic(self):
        """DOCX files are ZIP archives; first two bytes are PK."""
        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        assert result[:2] == b"PK"

    def test_docx_can_be_reparsed_by_python_docx(self):
        """python-docx should be able to open the generated bytes without error."""
        from docx import Document
        import io

        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_docx_contains_question_text(self):
        """Verify the stems appear somewhere in the document paragraphs."""
        from docx import Document
        import io

        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "labeled" in all_text or "flowchart" in all_text.lower()

    def test_docx_contains_table_for_each_mc(self):
        from docx import Document
        import io

        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        doc = Document(io.BytesIO(result))
        # Two MC questions → two tables
        assert len(doc.tables) == 2

    def test_docx_tables_have_five_rows(self):
        """Header row + 4 choice rows = 5 rows for a 4-choice question."""
        from docx import Document
        import io

        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        doc = Document(io.BytesIO(result))
        q1_table = doc.tables[0]
        assert len(q1_table.rows) == 5  # 1 header + 4 choices

    def test_docx_tables_have_four_columns(self):
        from docx import Document
        import io

        quiz = _packaged_quiz(_per_choice_payload())
        result = CorrectionDocRenderer().render_docx(quiz)
        doc = Document(io.BytesIO(result))
        for table in doc.tables:
            assert len(table.columns) == 4

    def test_docx_empty_rationales_no_crash(self):
        payload = {
            "version": "3.0-json",
            "title": "Empty",
            "items": [
                {"id": "q1", "type": "MC", "prompt": "Q?",
                 "choices": [{"id": "A", "text": "A", "correct": True}]}
            ],
            "rationales": [],
        }
        quiz = _packaged_quiz(payload)
        result = CorrectionDocRenderer().render_docx(quiz)
        assert isinstance(result, bytes)
        assert len(result) > 0
