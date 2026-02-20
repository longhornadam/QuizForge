"""Physical quiz DOCX handler.

Converts validated Quiz objects into printable DOCX files.
Generates student quiz + answer key + rationale sheet.
Performs NO validation - trusts input is perfect.
"""

from pathlib import Path
from typing import Dict, List
import re
from engine.core.quiz import Quiz
from engine.packaging.folder_creator import sanitize_filename
from engine.rendering.physical.styles.default_styles import (
    DOCX_FONT_FAMILY,
    DOCX_TITLE_SIZE,
    DOCX_HEADING_SIZE,
    DOCX_BODY_SIZE,
    DOCX_SMALL_SIZE,
    DOCX_MARGIN_TOP,
    DOCX_MARGIN_BOTTOM,
    DOCX_MARGIN_LEFT,
    DOCX_MARGIN_RIGHT,
    DOCX_PARA_SPACING_BEFORE,
    DOCX_PARA_SPACING_AFTER,
    DOCX_LINE_SPACING,
    MC_TWO_COLUMN_THRESHOLD,
    TABLE_HEADER_BOLD,
    TABLE_GRID_LINES,
    TABLE_COL_WIDTHS,
    LOG_ANSWER_CHOICE_STATS,
    LOG_QUESTION_LENGTH_STATS,
    LOG_POINT_VALUE_DISTRIBUTION,
    LOG_STIMULUS_USAGE
)


class PhysicalHandler:
    """Generate printable DOCX files."""
    
    def package(self, quiz: Quiz, output_base: str) -> Dict[str, str]:
        """Create student quiz, answer key, and rationale sheet DOCX files.
        
        Args:
            quiz: Validated Quiz object
            output_base: Base directory for output files
            
        Returns:
            Dict with 'quiz_path', 'key_path', 'rationale_path'
        """
        try:
            results = generate_physical_outputs(quiz, output_base)
            return results
        except Exception as e:
            # Log error and return partial results
            error_log_path = Path(output_base) / "physical_error.log"
            error_log_path.write_text(f"Physical generation failed: {e}", encoding='utf-8')
            return {
                'quiz_path': '',
                'key_path': '',
                'rationale_path': '',
                'error_log': str(error_log_path)
            }


def generate_physical_outputs(quiz: Quiz, output_folder: str) -> Dict[str, str]:
    """
    Generate all three physical DOCX files.
    
    Args:
        quiz: Validated Quiz object
        output_folder: Path to quiz-specific output folder
        
    Returns:
        {
            'quiz_path': str,
            'key_path': str, 
            'rationale_path': str,
            'log_path': str
        }
    """
    # Create student quiz DOCX
    quiz_path = _create_student_quiz(quiz, output_folder)
    
    # Create answer key DOCX
    key_path = _create_answer_key(quiz, output_folder)
    
    # Create rationale sheet DOCX
    rationale_path = _create_rationale_sheet(quiz, output_folder)
    
    # Generate validation log
    log_path = Path(output_folder) / "physical_validation.log"
    _log_validation_stats(quiz, str(log_path))
    
    # Return paths
    return {
        'quiz_path': quiz_path,
        'key_path': key_path,
        'rationale_path': rationale_path,
        'log_path': str(log_path)
    }


def _create_student_quiz(quiz: Quiz, output_folder: str) -> str:
    """Generate student-facing quiz DOCX."""
    from docx import Document
    from docx.shared import Inches, Pt
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(DOCX_MARGIN_TOP)
        section.bottom_margin = Inches(DOCX_MARGIN_BOTTOM)
        section.left_margin = Inches(DOCX_MARGIN_LEFT)
        section.right_margin = Inches(DOCX_MARGIN_RIGHT)
    
    # Name line — above the title to leave room for handwriting
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(f'Name: {"_" * 50}')
    name_run.font.name = DOCX_FONT_FAMILY
    name_run.font.size = Pt(DOCX_SMALL_SIZE)

    # Title line
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(quiz.title)
    title_run.bold = True
    title_run.font.name = DOCX_FONT_FAMILY
    title_run.font.size = Pt(DOCX_TITLE_SIZE)
    doc.add_paragraph()  # blank line
    
    # Process questions grouped by stimulus
    stimulus_groups = _group_questions_by_stimulus(quiz.questions)
    
    question_number = 1
    for group in stimulus_groups:
        if group['stimulus']:
            _render_stimulus_boxed(
                group['stimulus'], doc,
                stimulus_format=group.get('stimulus_format', 'text'),
                title=group.get('stimulus_title', ''),
                author=group.get('stimulus_author', ''),
            )
            doc.add_paragraph()  # spacing
        
        # Add questions in this group
        for q in group['questions']:
            _add_question_to_doc(doc, q, question_number)
            question_number += 1
    
    # Save document
    output_path = Path(output_folder) / f"{sanitize_filename(quiz.title)}.docx"
    doc.save(str(output_path))
    return str(output_path)


def _create_answer_key(quiz: Quiz, output_folder: str) -> str:
    """Generate teacher answer key DOCX."""
    from docx import Document
    from docx.shared import Inches, Pt
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(DOCX_MARGIN_TOP)
        section.bottom_margin = Inches(DOCX_MARGIN_BOTTOM)
        section.left_margin = Inches(DOCX_MARGIN_LEFT)
        section.right_margin = Inches(DOCX_MARGIN_RIGHT)
    
    # Add title
    title = doc.add_paragraph(f"{quiz.title} - Answer Key")
    title_format = title.runs[0].font
    title_format.name = DOCX_FONT_FAMILY
    title_format.size = Pt(DOCX_TITLE_SIZE)
    title_format.bold = True
    
    doc.add_paragraph()  # spacing
    
    # Get scorable questions only
    scorable_questions = quiz.scorable_questions()
    num_questions = len(scorable_questions)
    
    if num_questions > 0:
        # Create table: header + data rows
        table = doc.add_table(rows=num_questions + 1, cols=3)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(TABLE_COL_WIDTHS['question_num'])
        table.columns[1].width = Inches(TABLE_COL_WIDTHS['correct_answer'])
        table.columns[2].width = Inches(TABLE_COL_WIDTHS['points'])
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Question #'
        header_cells[1].text = 'Correct Answer'
        header_cells[2].text = 'Points'
        
        # Bold headers if configured
        if TABLE_HEADER_BOLD:
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = DOCX_FONT_FAMILY
                        run.font.size = Pt(DOCX_BODY_SIZE)
        
        # Data rows
        question_number = 1
        for idx, question in enumerate(quiz.questions, start=1):
            from engine.core.questions import StimulusItem, StimulusEnd
            
            # Skip non-scorable questions
            if isinstance(question, (StimulusItem, StimulusEnd)):
                continue
                
            row_cells = table.rows[question_number].cells
            
            # Question number
            row_cells[0].text = str(question_number)
            
            # Correct answer
            answer_text = _get_correct_answer_text(question)
            row_cells[1].text = answer_text
            
            # Points
            row_cells[2].text = str(question.points)
            
            # Style the data cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = DOCX_FONT_FAMILY
                        run.font.size = Pt(DOCX_BODY_SIZE)
            
            question_number += 1
        
        # Add total points
        doc.add_paragraph()
        total_points = quiz.total_points()
        total_para = doc.add_paragraph(f"Total Points: {total_points}")
        total_format = total_para.runs[0].font
        total_format.name = DOCX_FONT_FAMILY
        total_format.size = Pt(DOCX_BODY_SIZE)
        total_format.bold = True
    
    # Save document
    output_path = Path(output_folder) / f"{sanitize_filename(quiz.title)}_KEY.docx"
    doc.save(str(output_path))
    return str(output_path)


def _create_rationale_sheet(quiz: Quiz, output_folder: str) -> str:
    """Generate student rationale/corrections sheet DOCX."""
    from docx import Document
    from docx.shared import Inches, Pt
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(DOCX_MARGIN_TOP)
        section.bottom_margin = Inches(DOCX_MARGIN_BOTTOM)
        section.left_margin = Inches(DOCX_MARGIN_LEFT)
        section.right_margin = Inches(DOCX_MARGIN_RIGHT)
    
    # Add title
    title = doc.add_paragraph(f"{quiz.title} - Rationales")
    title_format = title.runs[0].font
    title_format.name = DOCX_FONT_FAMILY
    title_format.size = Pt(DOCX_TITLE_SIZE)
    title_format.bold = True
    
    # Add instructions
    instructions = doc.add_paragraph("Use this sheet to write corrections for questions you missed.")
    instructions_format = instructions.runs[0].font
    instructions_format.name = DOCX_FONT_FAMILY
    instructions_format.size = Pt(DOCX_BODY_SIZE)
    
    doc.add_paragraph()  # blank line
    
    # Build rationale lookup by item_id if available
    rationale_lookup = {}
    for entry in quiz.rationales or []:
        if isinstance(entry, dict):
            item_id = entry.get("item_id")
            if item_id:
                rationale_lookup[item_id] = entry.get("text") or entry.get("correct") or ""
        elif isinstance(entry, str):
            # fallback parsing "id: text"
            if ":" in entry:
                item_id, text = entry.split(":", 1)
                rationale_lookup[item_id.strip()] = text.strip()
    
    # Process questions
    question_number = 1
    for question in quiz.questions:
        from engine.core.questions import StimulusItem, StimulusEnd
        
        # Skip non-scorable questions
        if isinstance(question, (StimulusItem, StimulusEnd)):
            continue
        
        # Skip ESSAY and FILEUPLOAD (no rationales needed)
        if question.qtype in ['ESSAY', 'FILEUPLOAD']:
            question_number += 1
            continue
        
        rationale = None
        q_id = getattr(question, "forced_ident", None)
        if q_id and q_id in rationale_lookup:
            rationale = rationale_lookup[q_id]

        # Fallback to basic rationale if missing
        if not rationale:
            rationale = _format_rationale_with_answer(question)
        
        # Format: "Q# - Rationale text"
        p = doc.add_paragraph()
        q_prefix = p.add_run(f"Q{question_number} - ")
        q_prefix.bold = True
        q_prefix.font.name = DOCX_FONT_FAMILY
        q_prefix.font.size = Pt(DOCX_BODY_SIZE)
        
        rationale_run = p.add_run(rationale)
        rationale_run.font.name = DOCX_FONT_FAMILY
        rationale_run.font.size = Pt(DOCX_BODY_SIZE)
        
        # No blank lines needed between rationales
        
        question_number += 1
    

    # Footer removed as per updated requirements
    
    # Save document
    output_path = Path(output_folder) / f"{sanitize_filename(quiz.title)}_RATIONALE.docx"
    doc.save(str(output_path))
    return str(output_path)


def _log_validation_stats(quiz: Quiz, log_path: str) -> None:
    """Write validation statistics to log file."""
    with open(log_path, 'w', encoding='utf-8') as log:
        log.write("\n" + "="*60 + "\n")
        log.write("PHYSICAL QUIZ VALIDATION STATISTICS\n")
        log.write("="*60 + "\n\n")
        
        if LOG_ANSWER_CHOICE_STATS:
            _log_answer_choice_analysis(quiz, log)
        
        if LOG_QUESTION_LENGTH_STATS:
            _log_question_length_analysis(quiz, log)
        
        if LOG_POINT_VALUE_DISTRIBUTION:
            _log_point_distribution(quiz, log)
        
        if LOG_STIMULUS_USAGE:
            _log_stimulus_analysis(quiz, log)


def _log_answer_choice_analysis(quiz: Quiz, log):
    """Log MC answer choice statistics."""
    log.write("ANSWER CHOICE STATISTICS\n")
    log.write("-" * 40 + "\n")
    
    from engine.core.questions import MCQuestion
    mc_questions = [q for q in quiz.questions if isinstance(q, MCQuestion)]
    
    if not mc_questions:
        log.write("No multiple choice questions found.\n\n")
        return
    
    # Choice length stats
    all_choice_lengths = []
    for q in mc_questions:
        for choice in q.choices:
            all_choice_lengths.append(len(choice.text))
    
    log.write(f"Total MC questions: {len(mc_questions)}\n")
    log.write(f"Average choice length: {sum(all_choice_lengths)/len(all_choice_lengths):.1f} chars\n")
    log.write(f"Min choice length: {min(all_choice_lengths)} chars\n")
    log.write(f"Max choice length: {max(all_choice_lengths)} chars\n")
    
    # Correct answer distribution
    correct_positions = []
    for q in mc_questions:
        for idx, choice in enumerate(q.choices):
            if choice.correct:
                correct_positions.append(idx)
                break
    
    log.write(f"\nCorrect answer distribution:\n")
    for pos in range(max(correct_positions) + 1):
        count = correct_positions.count(pos)
        letter = chr(65 + pos)
        log.write(f"  {letter}: {count} ({count/len(mc_questions)*100:.1f}%)\n")
    
    # Two-column layout eligibility
    two_col_eligible = 0
    for q in mc_questions:
        choice_lengths = [len(choice.text) for choice in q.choices]
        if all(length < MC_TWO_COLUMN_THRESHOLD for length in choice_lengths):
            two_col_eligible += 1
    log.write(f"\nQuestions eligible for 2-column layout: {two_col_eligible}/{len(mc_questions)}\n")
    log.write("\n")


def _log_question_length_analysis(quiz: Quiz, log):
    """Log question text length statistics."""
    log.write("QUESTION LENGTH STATISTICS\n")
    log.write("-" * 40 + "\n")
    
    question_lengths = [len(q.prompt) for q in quiz.questions]
    
    log.write(f"Total questions: {len(quiz.questions)}\n")
    log.write(f"Average question length: {sum(question_lengths)/len(question_lengths):.1f} chars\n")
    log.write(f"Min question length: {min(question_lengths)} chars\n")
    log.write(f"Max question length: {max(question_lengths)} chars\n")
    
    # Identify very long questions
    long_threshold = 200
    long_questions = [(i+1, q) for i, q in enumerate(quiz.questions) if len(q.prompt) > long_threshold]
    
    if long_questions:
        log.write(f"\nQuestions over {long_threshold} chars (may need formatting review):\n")
        for q_num, q in long_questions:
            log.write(f"  Q{q_num}: {len(q.prompt)} chars\n")
    
    log.write("\n")


def _log_point_distribution(quiz: Quiz, log):
    """Log point value distribution."""
    log.write("POINT VALUE DISTRIBUTION\n")
    log.write("-" * 40 + "\n")
    
    from engine.core.questions import StimulusItem, StimulusEnd
    scorable_questions = [q for q in quiz.questions if not isinstance(q, (StimulusItem, StimulusEnd))]
    point_values = [q.points for q in scorable_questions]
    total_points = sum(point_values)
    
    log.write(f"Total points: {total_points}\n")
    log.write(f"Average points per question: {total_points/len(scorable_questions):.2f}\n")
    
    # Point value breakdown
    unique_values = sorted(set(point_values))
    log.write(f"\nPoint value breakdown:\n")
    for val in unique_values:
        count = point_values.count(val)
        log.write(f"  {val} pt: {count} questions ({count/len(scorable_questions)*100:.1f}%)\n")
    
    log.write("\n")


def _log_stimulus_analysis(quiz: Quiz, log):
    """Log stimulus passage usage statistics."""
    log.write("STIMULUS USAGE STATISTICS\n")
    log.write("-" * 40 + "\n")
    
    from engine.core.questions import StimulusItem, StimulusEnd
    questions_with_stimuli = []
    stimulus_groups = {}
    
    for i, q in enumerate(quiz.questions):
        if isinstance(q, StimulusItem):
            # Start of stimulus group
            current_stimulus = q.prompt
            if current_stimulus not in stimulus_groups:
                stimulus_groups[current_stimulus] = []
        elif isinstance(q, StimulusEnd):
            # End of stimulus group
            current_stimulus = None
        elif current_stimulus:
            # Question belongs to current stimulus
            questions_with_stimuli.append(q)
            stimulus_groups[current_stimulus].append(i + 1)  # 1-indexed
    
    log.write(f"Questions with stimuli: {len(questions_with_stimuli)}/{len(quiz.scorable_questions())}\n")
    
    if questions_with_stimuli:
        stimulus_lengths = [len(stimulus) for stimulus in stimulus_groups.keys()]
        log.write(f"Average stimulus length: {sum(stimulus_lengths)/len(stimulus_lengths):.1f} chars\n")
        log.write(f"Unique stimuli: {len(stimulus_groups)}\n")
        log.write(f"Questions per stimulus (avg): {len(questions_with_stimuli)/len(stimulus_groups):.1f}\n")
    
    log.write("\n")


def _group_questions_by_stimulus(questions):
    """Group questions by their stimulus associations."""
    from engine.core.questions import StimulusItem, StimulusEnd
    
    groups = []
    current_stimulus = None
    current_stimulus_format = None
    current_questions = []
    
    current_stimulus_title = ''
    current_stimulus_author = ''

    for q in questions:
        if isinstance(q, StimulusItem):
            # Save previous group if exists
            if current_questions:
                groups.append({
                    'stimulus': current_stimulus,
                    'stimulus_format': current_stimulus_format,
                    'stimulus_title': current_stimulus_title,
                    'stimulus_author': current_stimulus_author,
                    'questions': current_questions
                })
            # Start new stimulus group
            current_stimulus = q.prompt
            current_stimulus_format = _detect_stimulus_format(q.prompt)
            current_stimulus_title = getattr(q, 'title', '') or ''
            current_stimulus_author = getattr(q, 'author', '') or ''
            current_questions = []
        elif isinstance(q, StimulusEnd):
            # End current stimulus group
            if current_questions:
                groups.append({
                    'stimulus': current_stimulus,
                    'stimulus_format': current_stimulus_format,
                    'stimulus_title': current_stimulus_title,
                    'stimulus_author': current_stimulus_author,
                    'questions': current_questions
                })
            current_stimulus = None
            current_stimulus_format = None
            current_stimulus_title = ''
            current_stimulus_author = ''
            current_questions = []
        else:
            # Regular question
            current_questions.append(q)
    
    # Add final group
    if current_questions:
        groups.append({
            'stimulus': current_stimulus,
            'stimulus_format': current_stimulus_format,
            'stimulus_title': current_stimulus_title,
            'stimulus_author': current_stimulus_author,
            'questions': current_questions
        })
    
    return groups


def _add_question_to_doc(doc, question, question_number):
    """Add single question with smart formatting."""
    from docx.shared import Pt, Inches
    from engine.core.questions import (
        MCQuestion,
        TFQuestion,
        NumericalQuestion,
        MAQuestion,
        MatchingQuestion,
        FITBQuestion,
        OrderingQuestion,
        CategorizationQuestion,
        EssayQuestion,
        FileUploadQuestion,
    )
    
    # Add question text (render HTML with question number prefix)
    q_text = _clean_prompt_text(question.prompt)
    if isinstance(question, FITBQuestion):
        token = getattr(question, "blank_token", "")
        if token and f"[{token}]" in q_text:
            q_text = q_text.replace(f"[{token}]", "__________________")
    
    # For True/False questions, add a small label line above the stem
    if isinstance(question, TFQuestion):
        lbl_para = doc.add_paragraph()
        lbl_para.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
        lbl_para.paragraph_format.space_after = Pt(0)
        lbl_run = lbl_para.add_run("(True/False)")
        lbl_run.italic = True
        lbl_run.font.name = DOCX_FONT_FAMILY
        lbl_run.font.size = Pt(DOCX_SMALL_SIZE)

    # Create first paragraph with question number prefix
    p = doc.add_paragraph()
    prefix_run = p.add_run(f"{question_number}. ")
    prefix_run.font.name = DOCX_FONT_FAMILY
    prefix_run.font.size = Pt(DOCX_BODY_SIZE)
    
    # Parse HTML and append to this paragraph
    from html.parser import HTMLParser as _HTMLParser
    class QuestionHTMLParser(_HTMLParser):
        def __init__(self, para):
            super().__init__()
            self.para = para
            self.formatting_stack = []
        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag in ('strong', 'b'):
                self.formatting_stack.append(('bold', True))
            elif tag in ('em', 'i'):
                self.formatting_stack.append(('italic', True))
        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in ('strong', 'b'):
                self.formatting_stack = [f for f in self.formatting_stack if f[0] != 'bold']
            elif tag in ('em', 'i'):
                self.formatting_stack = [f for f in self.formatting_stack if f[0] != 'italic']
        def handle_data(self, data):
            if not data.strip():  # skip whitespace-only nodes (newlines between tags)
                return
            is_bold = any(fmt[0] == 'bold' for fmt in self.formatting_stack)
            is_italic = any(fmt[0] == 'italic' for fmt in self.formatting_stack)
            run = self.para.add_run(data)  # use data directly to preserve spaces
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)
            run.bold = is_bold
            run.italic = is_italic
    
    try:
        parser = QuestionHTMLParser(p)
        parser.feed(q_text)
    except Exception:
        # Fallback: just add as plain text
        p.add_run(q_text).font.name = DOCX_FONT_FAMILY
    
    p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
    p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
    p.paragraph_format.line_spacing = DOCX_LINE_SPACING
    
    if isinstance(question, (MCQuestion, MAQuestion)):
        choices = [choice.text for choice in question.choices]
        _add_mc_single_column(doc, choices)
    
    elif isinstance(question, TFQuestion):
        _add_tf_options(doc)
    
    elif isinstance(question, MatchingQuestion):
        _add_matching_block(doc, question.pairs)
    
    elif isinstance(question, FITBQuestion):
        # Inline blank already in prompt; no extra line
        pass
    
    elif isinstance(question, OrderingQuestion):
        _add_ordering_block(doc, question.items)
    
    elif isinstance(question, CategorizationQuestion):
        _add_categorization_block(doc, question)
    
    elif isinstance(question, NumericalQuestion):
        doc.add_paragraph('Answer: _____________')
    
    elif isinstance(question, EssayQuestion):
        doc.add_paragraph('Answer: _____________')
    
    elif isinstance(question, FileUploadQuestion):
        # Instructions only; no answer line
        pass
    
    else:
        doc.add_paragraph('Answer: _____________')
    
    # Add spacing after question
    doc.add_paragraph()


def _add_mc_two_column(doc, choices):
    """Add MC choices in 2-column table format."""
    from docx.shared import Pt
    
    table = doc.add_table(rows=(len(choices) + 1) // 2, cols=2)
    table.style = 'Table Grid'
    
    for idx, choice in enumerate(choices):
        row = idx // 2
        col = idx % 2
        cell = table.rows[row].cells[col]
        cell.text = f"{chr(65+idx)}. {choice}"
        
        # Style the cell text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = DOCX_FONT_FAMILY
                run.font.size = Pt(DOCX_BODY_SIZE)


def _add_mc_single_column(doc, choices):
    """Add MC choices in single column."""
    from docx.shared import Pt, Inches
    from html.parser import HTMLParser as _HTMLParser
    
    for idx, choice in enumerate(choices):
        # Create paragraph with choice letter prefix
        p = doc.add_paragraph()
        prefix_run = p.add_run(f"{chr(65+idx)}. ")
        prefix_run.font.name = DOCX_FONT_FAMILY
        prefix_run.font.size = Pt(DOCX_BODY_SIZE)
        
        # Parse choice text for HTML formatting
        class ChoiceHTMLParser(_HTMLParser):
            def __init__(self, para):
                super().__init__()
                self.para = para
                self.formatting_stack = []
            def handle_starttag(self, tag, attrs):
                tag = tag.lower()
                if tag in ('strong', 'b'):
                    self.formatting_stack.append(('bold', True))
                elif tag in ('em', 'i'):
                    self.formatting_stack.append(('italic', True))
            def handle_endtag(self, tag):
                tag = tag.lower()
                if tag in ('strong', 'b'):
                    self.formatting_stack = [f for f in self.formatting_stack if f[0] != 'bold']
                elif tag in ('em', 'i'):
                    self.formatting_stack = [f for f in self.formatting_stack if f[0] != 'italic']
            def handle_data(self, data):
                if not data.strip():  # skip whitespace-only nodes (newlines between tags)
                    return
                is_bold = any(fmt[0] == 'bold' for fmt in self.formatting_stack)
                is_italic = any(fmt[0] == 'italic' for fmt in self.formatting_stack)
                run = self.para.add_run(data)  # use data directly to preserve spaces
                run.font.name = DOCX_FONT_FAMILY
                run.font.size = Pt(DOCX_BODY_SIZE)
                run.bold = is_bold
                run.italic = is_italic
        
        try:
            parser = ChoiceHTMLParser(p)
            parser.feed(choice)
        except Exception:
            # Fallback: just add as plain text
            p.add_run(choice).font.name = DOCX_FONT_FAMILY
        
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
        p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
        p.paragraph_format.line_spacing = DOCX_LINE_SPACING


def _add_matching_block(doc, pairs: List):
    """Render matching as blanks plus legend letters."""
    from docx.shared import Pt, Inches
    letters = [chr(65 + idx) for idx in range(len(pairs))]

    # Definitions with blanks
    for pair in pairs:
        p = doc.add_paragraph(f"_____: {pair.answer}")
        p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
        p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)

    # Legend for terms
    for letter, pair in zip(letters, pairs):
        p = doc.add_paragraph(f"{letter} – {pair.prompt}")
        p.paragraph_format.left_indent = Inches(0.15)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)


def _add_ordering_block(doc, items):
    """Render ordering items with blanks for numbering."""
    from docx.shared import Pt, Inches
    for text in items:
        p = doc.add_paragraph(f"_____: {text.text if hasattr(text, 'text') else text}")
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
        p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)


def _add_categorization_block(doc, question):
    """Render categorization with category hints and blanks."""
    from docx.shared import Pt, Inches
    categories = question.categories if hasattr(question, "categories") else []
    if categories:
        cat_line = "Categories: " + ", ".join(categories)
        p = doc.add_paragraph(cat_line)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)

    # Combine items and distractors in display order
    items = getattr(question, "items", [])
    distractors = getattr(question, "distractors", []) or []

    for entry in items:
        label = entry.item_text if hasattr(entry, "item_text") else entry.get("label", "")
        p = doc.add_paragraph(f"_____: {label}")
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
        p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)

    for label in distractors:
        p = doc.add_paragraph(f"_____: {label}")
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
        p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)


def _add_tf_options(doc):
    """Add True/False options."""
    from docx.shared import Pt, Inches
    
    # Add options in a simple format
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(DOCX_PARA_SPACING_BEFORE)
    p.paragraph_format.space_after = Pt(DOCX_PARA_SPACING_AFTER)
    p.paragraph_format.line_spacing = DOCX_LINE_SPACING
    
    run = p.add_run("A. True    B. False")
    run.font.name = DOCX_FONT_FAMILY
    run.font.size = Pt(DOCX_BODY_SIZE)


def _get_correct_answer_text(question):
    """Get the correct answer text for a question."""
    from engine.core.questions import (
        MCQuestion, 
        TFQuestion, 
        NumericalQuestion,
        MAQuestion,
        MatchingQuestion,
        FITBQuestion,
        EssayQuestion,
        FileUploadQuestion
    )
    
    if isinstance(question, MCQuestion):
        # Find correct choice letter
        for idx, choice in enumerate(question.choices):
            if choice.correct:
                return chr(65 + idx)  # A, B, C, D...
        return "?"
    
    elif isinstance(question, TFQuestion):
        return "True" if question.answer_true else "False"
    
    elif isinstance(question, NumericalQuestion):
        return str(question.answer.answer)
    
    elif isinstance(question, MAQuestion):
        # List all correct choices
        correct_letters = []
        for idx, choice in enumerate(question.choices):
            if choice.correct:
                correct_letters.append(chr(65 + idx))
        return ", ".join(correct_letters) if correct_letters else "?"
    
    elif isinstance(question, MatchingQuestion):
        # List pairs as "left→right" format
        pairs = []
        for pair in question.pairs:
            # Format: "term → match"
            pairs.append(f"{pair.prompt} → {pair.answer}")
        return "; ".join(pairs) if pairs else "?"
    
    elif isinstance(question, FITBQuestion):
        # List all accepted answers
        if hasattr(question, "accepted_answers") and question.accepted_answers:
            return ", ".join(question.accepted_answers)

        # Use variants lists produced by importer
        variants = getattr(question, "variants", []) or []
        per_blank = getattr(question, "variants_per_blank", []) or []

        if per_blank:
            parts = []
            for idx, group in enumerate(per_blank, 1):
                if group:
                    unique = list(dict.fromkeys(str(v) for v in group))
                    parts.append(f"Blank {idx}: " + " | ".join(unique))
            if parts:
                return "; ".join(parts)

        if variants:
            unique = list(dict.fromkeys(str(v) for v in variants))
            return " | ".join(unique)

        return "?"
    
    elif isinstance(question, (EssayQuestion, FileUploadQuestion)):
        # These are subjectively graded
        return "See rubric"
    
    else:
        # Catch-all for unknown types
        return "See rubric"


def _format_rationale_with_answer(question):
    """
    Ensure rationale has answer embedded properly.
    
    Good: "Sarah is the antagonist because she opposes the protagonist."
    Bad: "The antagonist is Sarah." (answer not embedded in explanation)
    """
    # For now, since quiz data doesn't have rationales, generate basic ones
    return _generate_basic_rationale(question)


def _generate_basic_rationale(question):
    """Generate simple rationale with answer embedded and no teacher deferrals."""
    from engine.core.questions import (
        MCQuestion,
        TFQuestion,
        NumericalQuestion,
        MAQuestion,
        FITBQuestion,
        MatchingQuestion,
        OrderingQuestion,
        CategorizationQuestion,
        EssayQuestion,
        FileUploadQuestion,
    )

    if isinstance(question, MCQuestion):
        for idx, choice in enumerate(question.choices):
            if choice.correct:
                correct_letter = chr(65 + idx)
                correct_text = choice.text
                return f"The correct answer is {correct_letter}: {correct_text}."
        return "The correct answer is listed in the answer key."

    if isinstance(question, MAQuestion):
        answers = []
        for idx, choice in enumerate(question.choices):
            if choice.correct:
                answers.append(f"{chr(65 + idx)} ({choice.text})")
        if answers:
            return "Correct answers: " + "; ".join(answers)
        return "Correct answers are indicated in the key."

    if isinstance(question, TFQuestion):
        answer = "True" if question.answer_true else "False"
        return f"The correct answer is {answer}."

    if isinstance(question, NumericalQuestion):
        correct_value = str(question.answer.answer)
        return f"The correct numeric value is {correct_value}."

    if isinstance(question, FITBQuestion):
        per_blank = getattr(question, "variants_per_blank", []) or []
        variants = getattr(question, "variants", []) or []
        if per_blank:
            parts = []
            for idx, group in enumerate(per_blank, 1):
                if group:
                    unique = list(dict.fromkeys(str(v) for v in group))
                    parts.append(f"Blank {idx}: " + " | ".join(unique))
            if parts:
                return "Accepted answers: " + "; ".join(parts)
        if variants:
            unique = list(dict.fromkeys(str(v) for v in variants))
            return "Accepted answers: " + " | ".join(unique)
        return "Accepted answers are listed in the key."

    if isinstance(question, MatchingQuestion):
        pairs = []
        for pair in question.pairs:
            pairs.append(f"{pair.prompt} -> {pair.answer}")
        if pairs:
            return "Correct matches: " + "; ".join(pairs)
        return "Correct matches are listed in the key."

    if isinstance(question, OrderingQuestion):
        items = getattr(question, "items", []) or []
        labels = [getattr(item, "text", str(item)) for item in items]
        if labels:
            return "Correct sequence: " + " > ".join(labels)
        return "Correct sequence is shown in the key."

    if isinstance(question, CategorizationQuestion):
        entries = []
        for entry in getattr(question, "items", []) or []:
            label = getattr(entry, "item_text", "") or entry.get("label", "")
            cat = getattr(entry, "category_name", "") or entry.get("category", "")
            if label and cat:
                entries.append(f"{label} -> {cat}")
        if entries:
            return "Categories: " + "; ".join(entries)
        return "Category placements are shown in the key."

    if isinstance(question, (EssayQuestion, FileUploadQuestion)):
        return "Open response; evaluate with the provided rubric."

    return "Correct answer details are provided in the answer key."


def _clean_prompt_text(prompt: str) -> str:
    """Strip markdown code fences and trim whitespace for DOCX display."""
    # Remove simple ``` fences
    cleaned = re.sub(r"```[a-zA-Z]*", "", prompt)
    cleaned = cleaned.replace("```", "")
    return cleaned.strip()


def _detect_stimulus_format(raw_prompt: str) -> str:
    """Detect stimulus format from raw prompt text before fence-stripping.
    
    Returns 'poetry' if a poetry/verse code fence is found, else 'text'.
    """
    if not raw_prompt:
        return 'text'
    if re.search(r'```\s*(poetry|verse)', raw_prompt, re.IGNORECASE):
        return 'poetry'
    return 'text'


def _render_stimulus_boxed(raw_prompt: str, doc, stimulus_format: str = 'text',
                           title: str = '', author: str = '') -> None:
    """Render a stimulus with an attribution header followed by indented content.

    Attribution (title + author) is printed before the passage in a slightly
    larger font.  No border box — whitespace provides the visual separation.
    Prose paragraphs are indented 1 tab (0.5").  Poetry uses a 2-tab text column
    with line numbers sitting at the 1-tab (0.5") position.
    """
    from docx.shared import Pt, Inches

    PROSE_INDENT = Inches(0.5)   # 1 default tab

    # --- Attribution header (before content) ---
    if title:
        t_para = doc.add_paragraph()
        t_para.paragraph_format.space_before = Pt(0)
        t_para.paragraph_format.space_after = Pt(2)
        t_run = t_para.add_run(f'\u201c{title}\u201d')
        t_run.bold = True
        t_run.font.name = DOCX_FONT_FAMILY
        t_run.font.size = Pt(DOCX_HEADING_SIZE)

    if author:
        a_para = doc.add_paragraph()
        a_para.paragraph_format.space_before = Pt(0)
        a_para.paragraph_format.space_after = Pt(4)
        a_run = a_para.add_run(author)
        a_run.italic = True
        a_run.font.name = DOCX_FONT_FAMILY
        a_run.font.size = Pt(DOCX_BODY_SIZE)

    # --- Stimulus content ---
    if stimulus_format == 'poetry':
        _render_poetry_to_doc(raw_prompt, doc)
    else:
        paras = _render_html_to_paragraph(
            _clean_prompt_text(raw_prompt), doc,
            apply_italic=True, number_paragraphs=True
        )
        # Apply 1-tab left indent to every prose paragraph
        for p in (paras or []):
            p.paragraph_format.left_indent = PROSE_INDENT


def _render_poetry_to_doc(raw_prompt: str, doc) -> None:
    """Render a poetry stimulus as individual line paragraphs.

    Layout: all poem text sits at 2 tabs (1.0").
    Every 5th line gets its number printed at 1 tab (0.5") in the left gutter,
    and a tab character advances to the 1.0" text column.
    """
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    TAB1 = 0.5   # inches — line-number column (1 default tab)
    TAB2 = 1.0   # inches — poem text column  (2 default tabs)

    def _set_tab_stop(para, inches):
        """Add a single left-aligned tab stop at the given inch position."""
        pPr = para._p.get_or_add_pPr()
        tabs_elem = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(int(inches * 1440)))
        tabs_elem.append(tab)
        pPr.append(tabs_elem)

    cleaned = _clean_prompt_text(raw_prompt)
    lines = [ln for ln in cleaned.splitlines() if ln.strip()]

    for line_num, line_text in enumerate(lines, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(TAB2)

        is_numbered = (line_num % 5 == 0)

        if is_numbered:
            # Hang back to the number column, tab forward to the text column
            p.paragraph_format.first_line_indent = -Inches(TAB1)
            _set_tab_stop(p, TAB2)

            num_run = p.add_run(str(line_num))
            num_run.font.name = DOCX_FONT_FAMILY
            num_run.font.size = Pt(DOCX_SMALL_SIZE)
            num_run.italic = False
            num_run.bold = False

            tab_run = p.add_run('\t')
            tab_run.font.name = DOCX_FONT_FAMILY
            tab_run.font.size = Pt(DOCX_BODY_SIZE)

        line_run = p.add_run(line_text.strip())
        line_run.italic = True
        line_run.font.name = DOCX_FONT_FAMILY
        line_run.font.size = Pt(DOCX_BODY_SIZE)


def _log_validation_stats(quiz: Quiz, log_path: str) -> None:
    """Write validation statistics to log file."""
    with open(log_path, 'a', encoding='utf-8') as log:
        log.write("\n" + "="*60 + "\n")
        log.write("PHYSICAL QUIZ VALIDATION STATISTICS\n")
        log.write("="*60 + "\n\n")
        
        if LOG_ANSWER_CHOICE_STATS:
            _log_answer_choice_analysis(quiz, log)
        
        if LOG_QUESTION_LENGTH_STATS:
            _log_question_length_analysis(quiz, log)
        
        if LOG_POINT_VALUE_DISTRIBUTION:
            _log_point_distribution(quiz, log)
        
        if LOG_STIMULUS_USAGE:
            _log_stimulus_analysis(quiz, log)


def _log_answer_choice_analysis(quiz: Quiz, log):
    """Log MC answer choice statistics."""
    log.write("ANSWER CHOICE STATISTICS\n")
    log.write("-" * 40 + "\n")
    
    from engine.core.questions import MCQuestion
    mc_questions = [q for q in quiz.questions if isinstance(q, MCQuestion)]
    
    if not mc_questions:
        log.write("No multiple choice questions found.\n\n")
        return
    
    # Choice length stats
    all_choice_lengths = []
    for q in mc_questions:
        for choice in q.choices:
            all_choice_lengths.append(len(choice.text))
    
    log.write(f"Total MC questions: {len(mc_questions)}\n")
    log.write(f"Average choice length: {sum(all_choice_lengths)/len(all_choice_lengths):.1f} chars\n")
    log.write(f"Min choice length: {min(all_choice_lengths)} chars\n")
    log.write(f"Max choice length: {max(all_choice_lengths)} chars\n")
    
    # Correct answer distribution
    correct_positions = []
    for q in mc_questions:
        for idx, choice in enumerate(q.choices):
            if choice.correct:
                correct_positions.append(idx)
                break
    
    log.write(f"\nCorrect answer distribution:\n")
    for pos in range(max(correct_positions) + 1):
        count = correct_positions.count(pos)
        letter = chr(65 + pos)
        log.write(f"  {letter}: {count} ({count/len(mc_questions)*100:.1f}%)\n")
    
    # Two-column layout eligibility
    two_col_eligible = 0
    for q in mc_questions:
        choice_lengths = [len(choice.text) for choice in q.choices]
        if all(length < MC_TWO_COLUMN_THRESHOLD for length in choice_lengths):
            two_col_eligible += 1
    log.write(f"\nQuestions eligible for 2-column layout: {two_col_eligible}/{len(mc_questions)}\n")
    log.write("\n")


def _log_question_length_analysis(quiz: Quiz, log):
    """Log question text length statistics."""
    log.write("QUESTION LENGTH STATISTICS\n")
    log.write("-" * 40 + "\n")
    
    question_lengths = [len(q.prompt) for q in quiz.questions]
    
    log.write(f"Total questions: {len(quiz.questions)}\n")
    log.write(f"Average question length: {sum(question_lengths)/len(question_lengths):.1f} chars\n")
    log.write(f"Min question length: {min(question_lengths)} chars\n")
    log.write(f"Max question length: {max(question_lengths)} chars\n")
    
    # Identify very long questions
    long_threshold = 200
    long_questions = [(i+1, q) for i, q in enumerate(quiz.questions) if len(q.prompt) > long_threshold]
    
    if long_questions:
        log.write(f"\nQuestions over {long_threshold} chars (may need formatting review):\n")
        for q_num, q in long_questions:
            log.write(f"  Q{q_num}: {len(q.prompt)} chars\n")
    
    log.write("\n")


def _log_point_distribution(quiz: Quiz, log):
    """Log point value distribution."""
    log.write("POINT VALUE DISTRIBUTION\n")
    log.write("-" * 40 + "\n")
    
    from engine.core.questions import StimulusItem, StimulusEnd
    scorable_questions = [q for q in quiz.questions if not isinstance(q, (StimulusItem, StimulusEnd))]
    point_values = [q.points for q in scorable_questions]
    total_points = sum(point_values)
    
    log.write(f"Total points: {total_points}\n")
    log.write(f"Average points per question: {total_points/len(scorable_questions):.2f}\n")
    
    # Point value breakdown
    unique_values = sorted(set(point_values))
    log.write(f"\nPoint value breakdown:\n")
    for val in unique_values:
        count = point_values.count(val)
        log.write(f"  {val} pt: {count} questions ({count/len(scorable_questions)*100:.1f}%)\n")
    
    log.write("\n")


def _log_stimulus_analysis(quiz: Quiz, log):
    """Log stimulus passage usage statistics."""
    log.write("STIMULUS USAGE STATISTICS\n")
    log.write("-" * 40 + "\n")
    
    from engine.core.questions import StimulusItem, StimulusEnd
    questions_with_stimuli = []
    stimulus_groups = {}
    current_stimulus = None
    
    for i, q in enumerate(quiz.questions):
        if isinstance(q, StimulusItem):
            # Start of stimulus group
            current_stimulus = q.prompt
            if current_stimulus not in stimulus_groups:
                stimulus_groups[current_stimulus] = []
        elif isinstance(q, StimulusEnd):
            # End of stimulus group
            current_stimulus = None
        elif current_stimulus:
            # Question belongs to current stimulus
            questions_with_stimuli.append(q)
            stimulus_groups[current_stimulus].append(i + 1)  # 1-indexed
    
    log.write(f"Questions with stimuli: {len(questions_with_stimuli)}/{len(quiz.scorable_questions())}\n")
    
    if questions_with_stimuli:
        stimulus_lengths = [len(stimulus) for stimulus in stimulus_groups.keys()]
        log.write(f"Average stimulus length: {sum(stimulus_lengths)/len(stimulus_lengths):.1f} chars\n")
        log.write(f"Unique stimuli: {len(stimulus_groups)}\n")
        log.write(f"Questions per stimulus (avg): {len(questions_with_stimuli)/len(stimulus_groups):.1f}\n")
    
    log.write("\n")


def _render_html_to_paragraph(html_string: str, doc, apply_italic: bool = False, number_paragraphs: bool = False) -> List:
    """Parse HTML and render as formatted DOCX paragraphs and runs.
    
    Converts <p>, <strong>, <b>, <em>, <i>, <br> tags into proper python-docx formatting.
    Any other tags are stripped and their text content is preserved as plain runs.
    
    Args:
        html_string: Raw HTML string (e.g., "<p><strong>Bold</strong></p>")
        doc: Document object to add paragraphs to
        apply_italic: If True, apply italic to all runs (for stimulus text)
        number_paragraphs: If True, add a small (n) number at the start of each <p> block
        
    Returns:
        List of paragraphs created (for further formatting by caller if needed)
    """
    from docx.shared import Pt
    from html.parser import HTMLParser
    
    if not html_string or not html_string.strip():
        return []
    
    class HTMLToDocxParser(HTMLParser):
        """Parse HTML into DOCX formatting."""
        
        def __init__(self):
            super().__init__()
            self.paragraphs = []
            self.current_para = None
            self.formatting_stack = []  # Stack of ('bold'|'italic', True) tuples
            self.para_count = 0
            self.needs_number = False
            
        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag == 'p':
                self.current_para = None
                self.para_count += 1
                self.needs_number = number_paragraphs
            elif tag in ('strong', 'b'):
                self.formatting_stack.append(('bold', True))
            elif tag in ('em', 'i'):
                self.formatting_stack.append(('italic', True))
            elif tag == 'br':
                self.current_para = None
                
        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in ('strong', 'b'):
                self.formatting_stack = [f for f in self.formatting_stack if f[0] != 'bold']
            elif tag in ('em', 'i'):
                self.formatting_stack = [f for f in self.formatting_stack if f[0] != 'italic']
                
        def handle_data(self, data):
            if not data.strip():  # skip whitespace-only nodes (newlines between tags)
                return
            
            # Ensure we have a current paragraph
            if self.current_para is None:
                self.current_para = doc.add_paragraph()
                self.paragraphs.append(self.current_para)
                
                # Paragraph number at start if requested
                if self.needs_number:
                    num_run = self.current_para.add_run(f'({self.para_count})\t')
                    num_run.font.name = DOCX_FONT_FAMILY
                    num_run.font.size = Pt(DOCX_SMALL_SIZE)
                    num_run.bold = True
                    num_run.italic = False
                    self.needs_number = False
            
            # Determine formatting from stack
            is_bold = any(fmt[0] == 'bold' for fmt in self.formatting_stack)
            is_italic = any(fmt[0] == 'italic' for fmt in self.formatting_stack)
            
            # Use data directly (not stripped) to preserve spaces between words
            run = self.current_para.add_run(data)
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)
            run.bold = is_bold
            run.italic = is_italic or apply_italic
    
    try:
        parser = HTMLToDocxParser()
        parser.feed(html_string)
        paragraphs_created = parser.paragraphs
    except Exception:
        # Fallback: treat as plain text
        p = doc.add_paragraph(html_string)
        for run in p.runs:
            run.font.name = DOCX_FONT_FAMILY
            run.font.size = Pt(DOCX_BODY_SIZE)
            if apply_italic:
                run.italic = True
        paragraphs_created = [p]
    
    return paragraphs_created
